import config_data as config  # 必须在最前：设置 HF_HUB_OFFLINE
import logging
import os
import re
import hashlib
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import easyocr
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_chroma import Chroma
from langchain_community.embeddings import DashScopeEmbeddings
import tiktoken

logger = logging.getLogger(__name__)

_ocr_reader = None


def _get_ocr_reader():
    global _ocr_reader
    if _ocr_reader is None:
        logger.info("初始化 EasyOCR（首次加载较慢）...")
        _ocr_reader = easyocr.Reader(["ch_sim", "en"], gpu=False)
        logger.info("EasyOCR 初始化完成")
    return _ocr_reader

def md5(text: str):
    return hashlib.md5(text.encode()).hexdigest()

# ============================
# ✅ 【新增】文本清洗预处理（核心）
# ============================
def clean_text(text: str) -> str:
    # 1. 替换多种换行、空格为标准格式
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")
    # 2. 去掉连续空行（保留最多1个换行）
    text = re.sub(r"\n{3,}", "\n\n", text)
    # 3. 去掉行首行尾空白
    lines = [line.strip() for line in text.split("\n")]
    # 4. 去掉空行、空白行
    lines = [line for line in lines if line]
    # 5. 去掉网址、邮箱（干扰信息）
    lines = [re.sub(r"http[s]?://\S+", "", line) for line in lines]
    # 只匹配真正的邮箱格式（user@domain.tld），避免误删含 @ 的密码/令牌
    lines = [re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", "", line) for line in lines]
    # 6. 去掉无意义特殊符号串（如 ----、####、~~~~ 等装饰线）
    lines = [re.sub(r"^[-=~*#_]{3,}", "", line) for line in lines]
    # 7. 合并成干净文本
    cleaned = "\n".join(lines)
    # 8. 合并多个空格为1个
    cleaned = re.sub(r" +", " ", cleaned)
    return cleaned.strip()

def load_txt(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    except (FileNotFoundError, UnicodeDecodeError) as e:
        logger.error("读取文件失败: %s", e)
        raise
    return clean_text(text)


def load_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        pages = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                pages.append(text)
        doc.close()
    except Exception:
        logger.exception("PDF 解析失败: %s", file_path)
        raise
    return clean_text("\n".join(pages))


def load_docx(file_path):
    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
    except Exception:
        logger.exception("DOCX 解析失败: %s", file_path)
        raise
    return clean_text("\n".join(paragraphs))


def load_image(file_path):
    try:
        reader = _get_ocr_reader()
        result = reader.readtext(file_path)
        text = " ".join(item[1] for item in result)
    except Exception:
        logger.exception("OCR 识别失败: %s", file_path)
        raise
    return clean_text(text)


def load_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        return load_txt(file_path)
    elif ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".bmp"):
        return load_image(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")


def split_text(text):
    # 加载 token 编码器（通义/千问/GPT 通用）
    tokenizer = tiktoken.get_encoding("cl100k_base")

    # 计算 token 数量
    def tiktoken_len(txt):
        return len(tokenizer.encode(txt))

    # 分块器：使用 token 长度计算
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.chunk_size,       # 现在代表 token 数量
        chunk_overlap=config.chunk_overlap, # 现在代表 token 数量
        length_function=tiktoken_len,       # 关键：按 token 算长度
        separators=["\n\n", "\n", "。", "！", "？", "，", " "]
    )
    return splitter.split_text(text)

def ingest_file(file_path, embedding=None):
    try:
        text = load_document(file_path)
        chunks = split_text(text)
    except Exception:
        logger.exception("文本处理失败: %s", file_path)
        raise

    docs = []
    for chunk in chunks:
        docs.append(
            Document(
                page_content=chunk,
                metadata={"source": file_path, "id": md5(chunk)}
            )
        )

    try:
        if embedding is None:
            embedding = DashScopeEmbeddings(
                model=config.embedding_model_name,
                dashscope_api_key=config.DASHSCOPE_API_KEY,
            )

        db = Chroma(
            collection_name=config.collection_name,
            embedding_function=embedding,
            persist_directory=config.persist_directory,
        )

        # 去重：如果同一来源文件已入库，先删除旧文档
        existing = db.get(where={"source": file_path})
        if existing["ids"]:
            db.delete(ids=existing["ids"])
            logger.info("已清理旧版本：%s", file_path)

        db.add_documents(docs)
        logger.info("已清洗并入库：%s（%d 个分块）", file_path, len(docs))
    except Exception:
        logger.exception("入库失败: %s", file_path)
        raise

